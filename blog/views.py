from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from .models import Post
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.models import User
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
import docx
from .forms import IdeaCaptureForm, DocCapture
from django.forms import formset_factory

# Create your views here.
def home(request):
    context = {
        'posts': Post.objects.all()
    }
    return render(request, 'blog/home.html', context)


def about(request):
    #return HttpResponse('<h>Blog Home</h1>') # simple http response
    return render(request, 'blog/about.html', {'title': 'About'})  # calls the html file

def template(request):
    #return HttpResponse('<h>Blog Home</h1>') # simple http response
    doc = docx.Document()

    context = {'title': 'template'}
    if request.method == "POST":
        context['txt'] = request.POST["textfield"]
        testtxt = context['txt'] = request.POST["textfield"]
        doc.add_paragraph(testtxt)
        doc.save('C:/Users/John/Downloads/testa.docx')

    return render(request, 'blog/template.html', context)


def idea_capture_form(request):
    context = {'title': 'Idea Capture'}

    doc = docx.Document()
    context['form'] = IdeaCaptureForm()

    form = IdeaCaptureForm()
    #label = 'empty'

    if request.method == "POST":
        form = IdeaCaptureForm(request.POST)
        if form.is_valid():

            for count, field in enumerate(IdeaCaptureForm.declared_fields.keys()):
                context['form'] = IdeaCaptureForm(request.POST)
                text = form.cleaned_data.get(field)
                head = IdeaCaptureForm.labels[count]
                doc.add_paragraph(head)
                doc.add_paragraph(text)
                doc.save('C:/Users/John/Downloads/testing.docx')


    return render(request, 'blog/idea_capture.html', context)

def doc_capture(request):
    context = {'title': 'Doc Capture'}
    context['my_in'] = 'my_in'

    form = DocCapture()
    context['form'] = DocCapture(extra=0)
    if request.method == "POST":
        form = DocCapture(request.POST)
        if form.is_valid():
            print ("valid")
            #extras = form.cleaned_data.get('fieldsNew')
            extras = request.POST['no_chapters']
            context['form'] = DocCapture(request.POST, extra=extras)


    return render(request, 'blog/doc_capture.html', context)


class PostListView(ListView): # class based view
    model = Post
    template_name = 'blog/home.html'
    context_object_name = 'posts'
    ordering = ['-data_posted']
    paginate_by = 1

class PostDetailView(DetailView): # class based view
    model = Post

class PostDeletelView(LoginRequiredMixin, UserPassesTestMixin, DeleteView): # class based view
    model = Post
    success_url = '/'

    def test_func(self):
        post = self.get_object()
        if self.request.user == post.author:
            return True
        return False

class PostCreateView(LoginRequiredMixin, CreateView): # class based view
    model = Post
    fields = ['title', 'content']

    def form_valid(self, form):
        form.instance.author = self.request.user
        return super().form_valid(form)


class PostUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView): # class based view
    model = Post
    fields = ['title', 'content']

    def form_valid(self, form):
        form.instance.author = self.request.user
        return super().form_valid(form)

    def test_func(self):
        post = self.get_object()
        if self.request.user == post.author:
            return True
        return False


class UserPostListView(ListView): # only displays blogs from selected user
    model = Post
    template_name = 'blog/user_posts.html'
    context_object_name = 'posts'
    ordering = ['-data_posted']
    paginate_by = 3

    def get_queryset(self):
        user = get_object_or_404(User, username=self.kwargs.get('username'))
        return Post.objects.filter(author=user).order_by('-data_posted')



